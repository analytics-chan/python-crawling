# 네이버 뉴스 본문 링크 가져와서 워드 문서에 저장하기

import requests
from bs4 import BeautifulSoup
import time

# 워드 문서 다루기
from docx import Document

# 1. 워드 생성하기
document = Document()

# 검색어 가져오기
keyword = input('검색어를 입력하세요 >>> ')

# 여러 페이지 추출
lastPage = int(input('몇 페이지까지 추출하시겠습니까?'))

page_num = 1

for i in range(1, lastPage * 10, 10):
    # page = document.add_paragraph(f'{page_num}페이지')
    # page.font_name = "돋움"

    response = requests.get(f"https://search.naver.com/search.naver?where=news&sm=tab_pge&query={keyword}&start={i}")
    html = response.text
    soup = BeautifulSoup(html, 'html.parser')

    articles = soup.select('div.info_group')

    for a in articles:
        links = a.select('a.info')

        if len(links) >= 2:
            url = links[1].attrs['href']

            response = requests.get(url)
            html = response.text
            soup = BeautifulSoup(html, 'html.parser')

            if 'entertain' in response.url:
                title = soup.select_one('.end_tit')
                content = soup.select_one('#articeBody')
            elif 'sports' in response.url:
                title = soup.select_one('h4.title')
                content = soup.select_one('#newsEndContents')

                divs = content.select('div')
                
                for div in divs:
                    div.decompose()

                paras = content.select('p')

                for para in paras:
                    para.decompose()
            else:
                title = soup.select_one('#title_area > span')
                content = soup.select_one('#dic_area')

            time.sleep(0.3)

            document.add_heading(title.text.strip(), level=0)
            document.add_paragraph(url)
            document.add_paragraph(content.text.strip())

            # c.font_name = '돋움'
        
    # page = document.add_paragraph(f'{page_num}페이지')
    # page.font_name = "돋움"


# # 2. 워드 데이터 추가하기
# document.add_heading('기사 제목', level=0)
# link = document.add_paragraph('기사 링크')
# content = document.add_paragraph('기사 본문')

# # https://doitgrow.com/42
# # 한글 폰트 적용
# link.font_name = '굴림'
# content.font_name = '굴림'

# # para = document.add_paragraph()
# run = para.add_run('기사 링크')


# 3. 워드 저장하기
# document.save(r'C:\Users\withbrother\Desktop\python_test2\chapter4\news.docx')
fileName = input('파일 이름을 입력하세요 >>> ')
document.save(rf'C:\Users\withbrother\Desktop\python_test2\chapter4\{fileName}.docx')