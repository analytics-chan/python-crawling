# 워드 문서 다루기
from docx import Document

# 1. 워드 생성하기
document = Document()

# 2. 워드 데이터 추가하기
document.add_heading('기사 제목', level=0)
link = document.add_paragraph('기사 링크')
content = document.add_paragraph('기사 본문')

# https://doitgrow.com/42
# 한글 폰트 적용
link.font_name = '굴림'
content.font_name = '굴림'

# para = document.add_paragraph()
# run = para.add_run('기사 링크')

# 3. 워드 저장하기
document.save('text.docx')